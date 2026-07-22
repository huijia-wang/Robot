from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "DYNAMIXEL_XL330-M288-T_单电机样例复现使用手册_2026-07-22.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20252B"
MUTED = "5F6B76"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
HEADER_FILL = "E8EEF5"
PALE_RED = "FCE8E6"
RED = "A61B1B"
PALE_GREEN = "EAF4EA"
GREEN = "2F6B3C"
PALE_GOLD = "FFF4CE"
GOLD = "7A5B00"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D5DCE3", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 11, MUTED, 0, 10),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name not in ("Subtitle",)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.font.color.rgb = rgb(INK)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.05


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DYNAMIXEL XL330-M288-T  |  单电机复现手册")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("2026-07-22  |  第 ")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MUTED)
    add_page_field(footer)
    run = footer.add_run(" 页")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MUTED)

    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = first_footer.add_run("版本 1.0  |  XL330-M288-T")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MUTED)


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE, accent=DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = rgb(accent)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")


def add_code_block(doc: Document, lines: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=100, bottom=100, start=130, end=130)
    p = cell.paragraphs[0]
    p.style = doc.styles["Code Block"]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(lines)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D5DCE3", size="4")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def create_numbering_instance(doc: Document) -> int:
    """Create a real decimal list that restarts at 1 for each procedure."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_steps(doc: Document, items: list[str]) -> None:
    num_id = create_numbering_instance(doc)
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        p_pr = paragraph._p.get_or_add_pPr()
        old_num_pr = p_pr.find(qn("w:numPr"))
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = rgb(NAVY)
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
        prevent_row_split(table.rows[-1])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.12
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.3)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    set_table_geometry(table, widths)
    set_table_borders(table)
    return table


def add_source(doc: Document, ref: str, label: str, url: str, note: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"[{ref}] {label} — ")
    run.bold = True
    add_hyperlink(p, url, url)
    if note:
        p.add_run(f"。{note}")


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("可复现工程手册  /  REPRODUCTION GUIDE")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(BLUE)
    doc.add_paragraph("DYNAMIXEL XL330-M288-T\n单电机样例复现使用手册", style="Title")
    doc.add_paragraph(
        "PC + U2D2 + 5V 外部电源 + Python / DYNAMIXEL SDK；覆盖接线、运动、反馈采集、故障判断与 MCU 迁移要点",
        style="Subtitle",
    )

    rows = [
        ["对象", "DYNAMIXEL XL330-M288-T（TTL 3 针）"],
        ["最小目标", "固定一台电机，完成小角度位置运动，并输出温度、转速、电流、电压等 CSV"],
        ["样例文件", "DYNAMIXEL_XL330_M288_sample/dynamixel_xl330_m288_sample.py"],
        ["版本", "1.0  |  2026-07-22  |  基于 ROBOTIS 官方资料核对"],
    ]
    table = add_table(doc, ["项目", "内容"], rows, [1800, 7560])
    for row in table.rows[1:]:
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    add_callout(
        doc,
        "先看这一条：绝对不要用 12V",
        "XL330-M288-T 的官方输入范围是 3.7–6.0V，推荐 5.0V。U2D2 不给电机供电；U2D2 Power Hub 也不能被当作把 12V 自动降成 5V 的电源。请给 Power Hub 输入已经稳压好的 5.0V，并在断电状态下接线。 [R1][R3][R4]",
        fill=PALE_RED,
        accent=RED,
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "复现完成的判定",
        "Wizard 能扫描到型号编号（Model Number）1200；Python 能 Ping 成功；电机只做受控小角度运动；CSV 中位置、速度、电流、输入电压、内部温度连续变化且 Hardware Error Status 为 0。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    props = doc.core_properties
    props.title = "DYNAMIXEL XL330-M288-T 单电机样例复现使用手册"
    props.subject = "XL330-M288-T, U2D2, TTL UART, Python, DYNAMIXEL SDK"
    props.author = "OpenAI Codex"
    props.keywords = "DYNAMIXEL, XL330-M288-T, U2D2, UART, TTL, Protocol 2.0"
    add_masthead(doc)

    doc.add_heading("1. 结论先行", level=1)
    doc.add_paragraph(
        "对这个型号，最稳妥的复现路径是先做“单电机桌面台架”，而不是直接接完整灵巧手：PC 负责发包和记录，U2D2 负责 USB 与半双工 TTL 转换，5V 外部电源负责电机功率。成功后再把同一套控制表地址和 Protocol 2.0 包迁移到 MCU。"
    )
    add_code_block(
        doc,
        "控制链：PC / Python  ->  USB  ->  U2D2  ->  3-pin TTL half-duplex  ->  XL330-M288-T\n"
        "供电链：regulated 5.0 V  ->  U2D2 Power Hub  ->  VDD / GND  ->  XL330-M288-T",
    )
    add_bullets(
        doc,
        [
            "电机：XL330-M288-T，18g，减速比 288.4:1，12bit 绝对编码器，4096 pulse/rev。",
            "信号：TTL 多点总线，3.3V 逻辑且兼容 5V；半双工异步串行，8 数据位、1 停止位、无校验（8N1）。",
            "协议：DYNAMIXEL Protocol 2.0。出厂默认通常是 ID=1、57,600bps、Position Control Mode(3)。",
            "反馈：位置、速度、电源侧输入电流、PWM、输入电压、内部温度、轨迹、运动状态、硬件错误。",
            "软件：先用 DYNAMIXEL Wizard 2.0 做发现/诊断，再用 Python + DYNAMIXEL SDK 运行本项目样例。",
        ],
    )

    doc.add_heading("2. XL330-M288-T 已确认规格", level=1)
    add_table(
        doc,
        ["项目", "官方规格", "复现时的含义"],
        [
            ["输入电压", "3.7–6.0V；推荐 5.0V", "使用稳压 5.0V；禁止 12V"],
            ["堵转转矩 @5V", "0.52N·m，1.47A", "堵转仅是瞬时性能点，不应长时间保持"],
            ["无负载速度 @5V", "103rev/min", "实测受电压、负载、Profile 限制影响"],
            ["其他电压点", "3.7V: 0.42N·m / 76rpm；6.0V: 0.60N·m / 123rpm", "供电改变会同步改变速度和转矩上限"],
            ["编码器", "非接触绝对编码器 AS5601，12bit，360°", "4096 pulse/rev；约 0.088°/pulse"],
            ["外形 / 重量", "20×34×26mm / 18g", "适合单指、关节或桌面演示"],
            ["温度范围", "-5–+70°C", "默认内部温度限制为 70°C；样例软件提前在 65°C 停止"],
            ["通信速率", "9,600bps–4Mbps", "复现先用默认 57,600bps，稳定后再提速"],
            ["型号编号", "Model Number = 1200", "样例 Ping 后强制核对，避免套错控制表"],
        ],
        [1900, 3000, 4460],
    )
    p = doc.add_paragraph()
    p.add_run("资料依据：").bold = True
    p.add_run("ROBOTIS XL330-M288-T 官方型号页与控制表 [R1]；日本商品页用于确认具体销售型号与同捆三针电缆 [R9]。")

    doc.add_heading("3. 信号、UART 与协议", level=1)
    doc.add_heading("3.1 这不是普通的 TX/RX 双线串口", level=2)
    doc.add_paragraph(
        "XL330 的 DATA 只有一根：控制器发送 Instruction Packet 后，必须释放总线并切换到接收方向，电机再返回 Status Packet。U2D2 已经替 PC 处理方向切换；若以后改为 MCU 裸机 UART，需要外部半双工收发/三态方向控制，并在最后一个停止位真正发完后才切到接收。Protocol 2.0 规定包内相邻字节间隔不能超过 1.5ms。 [R2]"
    )
    add_table(
        doc,
        ["层", "确认值", "说明"],
        [
            ["物理层", "TTL 多点总线；3.3V 逻辑，兼容 5V", "不是 RS-485，也不是独立 TX/RX"],
            ["串口格式", "半双工、异步、8N1", "8bit、1 stop、no parity"],
            ["链路速率", "默认 57,600bps；可设 9,600–4Mbps", "PC/U2D2 与电机必须一致"],
            ["协议", "DYNAMIXEL Protocol 2.0", "Header、ID、Length、Instruction、Parameters、CRC"],
            ["请求/响应", "Controller→Instruction；Motor→Status", "同一时刻只能一方驱动 DATA"],
            ["总线标识", "ID 0–252；广播 ID 254", "多电机时每台主 ID 必须唯一"],
        ],
        [1700, 2700, 4960],
    )

    doc.add_heading("3.2 三针定义", level=2)
    add_table(
        doc,
        ["Pin", "名称", "本样例连接", "注意"],
        [
            ["1", "GND", "5V 电源负端 / Power Hub GND", "PC、接口和电机应有可靠参考地"],
            ["2", "VDD", "稳压 5.0V", "只允许 3.7–6.0V；不要接 12V"],
            ["3", "DATA", "U2D2 的 3Pin TTL DATA", "半双工；不要接到 RS-485 口"],
        ],
        [900, 1300, 3400, 3760],
    )
    doc.add_paragraph(
        "电机侧连接器为 JST EHR-03，基板针座 B3B-EH-A。拔插前必须切断外部电源。 [R1]"
    )

    doc.add_heading("4. 采购清单与接线", level=1)
    add_table(
        doc,
        ["部件", "数量", "建议/作用", "是否必须"],
        [
            ["XL330-M288-T", "1", "本手册目标电机；TTL 3 针", "必须"],
            ["U2D2", "1", "PC USB ↔ DYNAMIXEL TTL/RS-485 转换器", "PC 路径必须"],
            ["U2D2 Power Hub Board", "1", "把外部电源和 U2D2/电机连接起来", "强烈建议"],
            ["Robot Cable-X3P", "1–2", "GND/VDD/DATA 三针电缆；商品通常附 180mm 一根", "必须"],
            ["稳压 5.0V 电源", "1", "工程建议单电机 ≥2A；推荐 5V/3A、带限流", "必须"],
            ["USB 数据线", "1", "连接 PC 与 U2D2；2025-08 后 U2D2 为 USB-C", "必须"],
            ["夹具/支架/舵盘", "1", "把电机固定并留出舵盘旋转空间", "必须"],
        ],
        [2400, 850, 4320, 1790],
    )
    add_callout(
        doc,
        "电源容量说明",
        "“5V/3A”是按官方 5V 堵转电流 1.47A 加启动/线损裕量给出的工程建议，不是 ROBOTIS 指定的唯一电源型号。请设置电流限制并避免堵转。U2D2 Power Hub 的板级范围虽为 3.5–24V、10A max，但给 XL330 时输入仍必须是约 5V。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.add_heading("4.1 断电接线步骤", level=2)
    add_steps(
        doc,
        [
            "关闭 Power Hub 开关，并断开 5V 电源输入；确认电源旋钮或输出设置为 5.0V。",
            "将 U2D2 固定在 Power Hub 上，用 3Pin TTL 线连接 U2D2 的 TTL 口与 Power Hub。",
            "用 Robot Cable-X3P 将 Power Hub 的 3Pin TTL 口连接到 XL330-M288-T。",
            "确认 Pin 1=GND、Pin 2=VDD、Pin 3=DATA；检查极性，确保没有接入 RS-485 4Pin 口。",
            "把电机固定在夹具上，舵盘周围清空；外部电源开关应随手可断。",
            "连接 PC USB；再接入稳压 5V。Power Hub 的多个电源输入只允许选一个。",
            "最后打开 Power Hub，先观察异味、发热、烟、异常声；异常时立即断电。",
        ],
    )

    doc.add_heading("5. 软件环境", level=1)
    add_table(
        doc,
        ["软件", "本手册基线", "用途"],
        [
            ["Windows", "Windows 10/11 64-bit", "Wizard 2.0 官方支持环境"],
            ["DYNAMIXEL Wizard 2.0", "使用当前官方安装包", "扫描、改 ID/波特率/模式、诊断、实时曲线、固件"],
            ["FTDI VCP 驱动", "U2D2 未出现 COM 口时安装", "让系统识别 U2D2"],
            ["Python", "建议 Python 3.10+；3.x", "运行样例与 CSV 记录"],
            ["dynamixel-sdk", "requirements.txt 固定为 4.0.5", "串口、Protocol 2.0、读写 API"],
        ],
        [2500, 3000, 3860],
    )
    doc.add_paragraph(
        "Wizard 2.0 支持 Windows 10/11 64-bit、Ubuntu 22.04/24.04 64-bit、macOS 13+；SDK 提供 C/C++/Python 等语言。 [R5][R6]"
    )

    doc.add_heading("6. 第一次上电：先用 Wizard 2.0", level=1)
    add_steps(
        doc,
        [
            "安装并打开 DYNAMIXEL Wizard 2.0；若看不到 U2D2 的 COM 口，安装 FTDI VCP 驱动。",
            "打开 Power Hub 的 5V 电源；在 Wizard 中选择 U2D2 的 COM 口。",
            "优先扫描 Protocol 2.0、57,600bps、ID 1。找不到时再扫描其他波特率/ID，不要盲目改线。",
            "确认检测结果为 XL330-M288-T，Model Number 为 1200；读取 Present Input Voltage 应约为 5.0V。",
            "确认 Operating Mode(11)=3（Position Control Mode），Torque Enable(64)=0。",
            "把 Goal Position(116) 设为当前 Present Position(132) 附近，再开 Torque，做约 10–20° 的小动作。",
            "观察 Present Temperature、Present Current、Present Velocity 和 Hardware Error Status；确认无异常后关闭 Torque。",
            "完全退出 Wizard。Wizard 占用 COM 口时，Python 不能同时打开同一端口。",
        ],
    )
    add_callout(
        doc,
        "不要在初次测试时做的事",
        "不要先改 ID/波特率；不要直接写极限位置；不要在未固定电机时开 Torque；不要让舵盘顶住机械限位；不要长时间堵转。",
        fill=PALE_RED,
        accent=RED,
    )

    doc.add_heading("7. Python 样例安装与运行", level=1)
    doc.add_paragraph(
        "样例目录为 DYNAMIXEL_XL330_M288_sample。程序先 Ping 并核对型号编号，再读取 Operating Mode；默认只做约 20° 的相对运动，持续采样并写 CSV，结束或 Ctrl+C 时尝试关闭 Torque。"
    )
    doc.add_heading("7.1 建立环境", level=2)
    add_code_block(
        doc,
        'cd "C:\\Users\\12192\\Documents\\新唐+灵巧手\\DYNAMIXEL_XL330_M288_sample"\n'
        "py -3 -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "python -m pip install -r requirements.txt",
    )
    doc.add_paragraph("PowerShell 禁止激活脚本时，可直接调用虚拟环境中的 python.exe：")
    add_code_block(
        doc,
        ".\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt\n"
        ".\\.venv\\Scripts\\python.exe dynamixel_xl330_m288_sample.py --self-test",
    )

    doc.add_heading("7.2 推荐运行顺序", level=2)
    doc.add_paragraph("A. 不接硬件也能做的离线自检：")
    add_code_block(doc, "python dynamixel_xl330_m288_sample.py --self-test")
    doc.add_paragraph("B. 只读取、不改变任何寄存器：")
    add_code_block(
        doc,
        "python dynamixel_xl330_m288_sample.py --port COM3 --read-only --duration 10",
    )
    doc.add_paragraph("C. 默认小动作并记录 5 秒：")
    add_code_block(doc, "python dynamixel_xl330_m288_sample.py --port COM3")
    doc.add_paragraph("D. 指定绝对目标位置：")
    add_code_block(
        doc,
        "python dynamixel_xl330_m288_sample.py --port COM3 --goal-pulse 2300 --duration 8",
    )
    doc.add_paragraph(
        "如果 Operating Mode 不是 3，程序会拒绝运动且不会默认改 EEPROM。优先在 Wizard 中改为 Position Control Mode；确认后也可显式使用 --set-position-mode。"
    )

    doc.add_heading("7.3 程序做了哪些安全处理", level=2)
    add_bullets(
        doc,
        [
            "Ping 返回型号编号不是 1200 时默认停止，避免拿 XL330 地址表驱动其他型号。",
            "不带 --set-position-mode 时不改 Operating Mode EEPROM；EEPROM 只能在 Torque=0 时写。",
            "开 Torque 前，先把 Goal Position 预装为当前单圈角度，降低突然跳转风险。",
            "默认 Profile Velocity=50，即约 11.45rpm；目标仅约 20°。",
            "温度 ≥65°C、供电不在 3.7–6.0V 或 Hardware Error 非零时，尝试关闭 Torque。",
            "无论正常结束、异常还是 Ctrl+C，finally 都会尝试 Torque Off 并关闭串口。",
            "通信已断时软件关断可能失败，所以物理电源开关仍是最终安全措施。",
        ],
    )

    doc.add_heading("8. 能读到的数据与换算", level=1)
    doc.add_paragraph(
        "下表来自 XL330-M288-T Control Table。多字节有符号量使用二进制补码；样例已经完成 16/32 位符号转换。 [R1]"
    )
    add_table(
        doc,
        ["参数", "地址/字节", "换算", "解释/用途"],
        [
            ["Hardware Error Status", "70 / 1", "bit field", "电压、过热、供电/电击、过载错误"],
            ["Realtime Tick", "120 / 2", "raw × 1ms", "0–32767，溢出后回到 0"],
            ["Moving", "122 / 1", "0/1", "速度超过阈值或 Profile 正在执行"],
            ["Moving Status", "123 / 1", "bit field", "Profile 类型、进行中、跟随误差、到位"],
            ["Present PWM", "124 / 2 signed", "约 raw × 0.113%", "当前逆变器 PWM；正负表示方向"],
            ["Present Current", "126 / 2 signed", "raw × 1mA", "XL330 的输入电源侧电流，不是相电流"],
            ["Present Velocity", "128 / 4 signed", "raw × 0.229rpm", "输出轴速度；正负表示方向"],
            ["Present Position", "132 / 4 signed", "raw × 360/4096°", "单圈 0–4095；特定模式下可为连续多圈值"],
            ["Velocity Trajectory", "136 / 4 signed", "raw × 0.229rpm", "Profile 生成的目标速度轨迹"],
            ["Position Trajectory", "140 / 4 signed", "raw × 360/4096°", "Profile 生成的目标位置轨迹"],
            ["Present Input Voltage", "144 / 2", "raw × 0.1V", "电机内部看到的供电电压"],
            ["Present Temperature", "146 / 1", "raw × 1°C", "电机内部温度，不是环境温度"],
        ],
        [2600, 1600, 2000, 3160],
    )

    doc.add_heading("8.1 CSV 列", level=2)
    doc.add_paragraph(
        "默认文件名 xl330_log.csv，使用 UTF-8 with BOM，Excel 可直接打开。主要列包括 timestamp_utc、elapsed_s、position_deg、velocity_rpm、current_mA、pwm_percent、input_voltage_V、temperature_C、moving、moving_status、hardware_error_text。"
    )
    add_callout(
        doc,
        "电流不是直接的“力”",
        "XL330 官方明确说明，它测量的是输入电源侧电流，与快速变化的直流电机相电流不同。即使 5V 堵转数据给出约 0.354N·m/A，也不能把任意动态电流直接当成精确轴转矩；摩擦、加速度、齿轮效率、温度、PWM 和机构臂长都会影响结果。若要得到指尖力，需要在目标机构上用力传感器/电子秤做标定。 [R1]",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.add_heading("9. 让电机运动的主要写入参数", level=1)
    add_table(
        doc,
        ["参数", "地址/字节", "单位/范围", "作用"],
        [
            ["Operating Mode", "11 / 1 EEPROM", "0,1,3,4,5,16", "选择电流、速度、单圈位置、多圈、限流位置或 PWM"],
            ["Torque Enable", "64 / 1 RAM", "0/1", "0=输出关闭；1=输出开启且锁定 EEPROM"],
            ["Goal PWM", "100 / 2 signed", "约 0.113%/unit", "PWM 模式直接控制；其他模式中作为输出限制"],
            ["Goal Current", "102 / 2 signed", "1mA/unit", "电流模式目标；限流位置模式中的电流上限"],
            ["Goal Velocity", "104 / 4 signed", "0.229rpm/unit", "速度模式目标"],
            ["Profile Acceleration", "108 / 4", "214.577rpm²/unit", "速度型 Profile 的加速度"],
            ["Profile Velocity", "112 / 4", "0.229rpm/unit", "位置模式的 Profile 最大速度"],
            ["Goal Position", "116 / 4 signed", "1 pulse", "单圈位置模式通常 0–4095"],
        ],
        [2500, 1800, 2400, 2660],
    )
    doc.add_heading("9.1 Operating Mode 对照", level=2)
    add_table(
        doc,
        ["值", "模式", "典型用途", "本手册"],
        [
            ["0", "Current Control", "夹持/电流控制", "后续实验；先标定"],
            ["1", "Velocity Control", "轮式连续旋转", "非首选"],
            ["3", "Position Control", "0–360° 单圈关节", "默认使用"],
            ["4", "Extended Position", "±256rev 多圈", "卷绕机构时再用"],
            ["5", "Current-based Position", "位置 + 电流限制", "灵巧手夹持后续重点"],
            ["16", "PWM Control", "直接 PWM 电压控制", "风险较高，非入门步骤"],
        ],
        [800, 2800, 3000, 2760],
    )

    doc.add_heading("10. 建议的复现实验", level=1)
    doc.add_heading("实验 A：静态通信与传感数据", level=2)
    add_bullets(
        doc,
        [
            "运行 --read-only 10 秒，确认电压约 5.0V、温度合理、Hardware Error=0。",
            "手动缓慢转动舵盘（Torque 必须关闭），观察 Position 是否按约 0.088°/pulse 变化。",
            "保存 CSV，作为接线和单位换算的基线记录。",
        ],
    )
    doc.add_heading("实验 B：20° 小角度位置动作", level=2)
    add_bullets(
        doc,
        [
            "固定电机并清空舵盘周围；运行默认命令。",
            "检查 position_deg 向目标变化、velocity_rpm 先升后降、Moving 最终回到 0。",
            "确认程序结束后 Torque Off；用手轻转舵盘应恢复可转动状态。",
        ],
    )
    doc.add_heading("实验 C：轻负载对比", level=2)
    add_bullets(
        doc,
        [
            "在舵盘加同一方向、可重复的小负载；保持相同目标、供电、Profile。",
            "比较空载/负载的峰值 current_mA、pwm_percent、到位时间和温升。",
            "不要堵转；若需要“力”结论，加入独立力传感器并做多点标定。",
        ],
    )
    add_table(
        doc,
        ["验收项", "合格判定", "记录"],
        [
            ["识别", "Model Number=1200；ID/baud 与设定一致", "截图或终端输出"],
            ["供电", "运行中 3.7–6.0V；建议约 5.0V 且无明显跌落", "CSV voltage"],
            ["位置", "方向正确、目标在限位内、无突然跳转", "CSV position"],
            ["速度", "曲线与动作一致，停止后接近 0rpm", "CSV velocity"],
            ["电流/PWM", "负载增加时趋势合理，无长期饱和", "CSV current/PWM"],
            ["温度", "低于 65°C 软件阈值；无持续快速上升", "CSV temperature"],
            ["错误", "Hardware Error Status=0", "CSV error"],
            ["结束状态", "程序退出后 Torque=0，串口释放", "手动确认"],
        ],
        [2100, 5200, 2060],
    )

    doc.add_heading("11. 故障排查", level=1)
    add_table(
        doc,
        ["现象", "优先检查", "处理"],
        [
            ["没有 COM 口", "USB 数据线、U2D2 指示灯、FTDI 驱动", "换数据线/USB 口；安装 VCP 驱动"],
            ["Wizard/Python 找不到电机", "外部 5V 是否开启；TTL 3Pin；ID/baud/protocol", "先回到 57,600bps、ID1、Protocol2 扫描"],
            ["串口被占用", "Wizard、串口工具、另一个 Python", "完全退出占用程序后重试"],
            ["能 Ping 但不能写 EEPROM", "Torque Enable 是否为 1", "先写 Torque=0，再改 Operating Mode/ID/baud"],
            ["一开 Torque 就跳动", "旧 Goal Position、Homing Offset、当前位置", "断电；重新固定；先 Goal=Present 再开 Torque"],
            ["电压低/反复重启", "电源电流、线损、接头、共地", "用更稳的 5V 电源和短粗线；不要靠 USB 给电机供电"],
            ["温度快速上升", "堵转、机构卡住、负载过大、Profile 太激进", "立即 Torque Off/断电，解除机械负载后冷却"],
            ["Hardware Error 0x01", "输入电压超限", "测量 VDD；恢复 3.7–6.0V 后重启"],
            ["Hardware Error 0x04", "内部过热", "断电冷却，降低负载/占空比"],
            ["Hardware Error 0x10", "电气冲击或供电不足", "检查电源容量、极性、接头和地"],
            ["Hardware Error 0x20", "持续过载", "解除卡滞/负载，重启后低速小动作验证"],
        ],
        [2500, 3400, 3460],
    )
    doc.add_paragraph(
        "出现 Shutdown 后，官方说明需要 Reboot 才能再次 Torque On；若 LED 持续闪烁，应先排除根因，不能只重复开 Torque。 [R1]"
    )

    doc.add_heading("12. 迁移到 MCU 的工程要点", level=1)
    doc.add_paragraph(
        "PC/U2D2 复现通过后，MCU 侧需要替代的是“USB 串口适配层”，电机控制表和 Protocol 2.0 语义保持不变。建议先保留单电机、5V 独立供电和小动作验收，再扩展多电机。"
    )
    add_table(
        doc,
        ["模块", "MCU 侧要求", "最小验证"],
        [
            ["UART", "57,600bps、8N1；发送完成以 shift-register empty 为准", "发 Ping 包，收到合法 Status 包"],
            ["半双工方向", "一根 DATA；TX 时驱动，发送完立即高阻/切 RX", "示波器/逻辑分析仪看不到总线争用"],
            ["电平/缓冲", "XL330 DATA 为 3.3V 逻辑且兼容 5V；推荐三态缓冲/方向控制", "确认 MCU 引脚电压与输入保护"],
            ["Protocol 2.0", "Header、Length、ID、Instruction、Byte Stuffing、CRC-16", "Ping、Read、Write 三种包单元测试"],
            ["时序", "包内 byte-to-byte ≤1.5ms；按 Return Delay 等待响应", "连续 1000 次读无 CRC/超时错误"],
            ["电源", "电机独立稳压 5V；与 MCU 共地；不要从 MCU 3.3V rail 供电", "动作时电压不跌出 3.7–6.0V"],
            ["软件安全", "超时、温度、电压、Hardware Error 触发 Torque Off", "模拟断线/过阈值并验证状态机"],
        ],
        [1800, 4900, 2660],
    )
    add_callout(
        doc,
        "推荐迁移顺序",
        "1) MCU 只做 Ping；2) 读 132/128/126/144/146；3) Torque Off 状态写 RAM 测试；4) 小角度 Position Mode；5) 加 Bus Watchdog；6) 再做多电机 Sync/Bulk Read。不要一开始就移植完整机械手或电流闭环。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    doc.add_heading("13. 文件说明", level=1)
    add_table(
        doc,
        ["文件", "用途", "验证状态"],
        [
            ["DYNAMIXEL_XL330_M288_sample/dynamixel_xl330_m288_sample.py", "单电机运动、监控、CSV、安全停止", "已通过 Python 语法检查和离线 self-test；未在实物上运行"],
            ["DYNAMIXEL_XL330_M288_sample/requirements.txt", "固定 dynamixel-sdk==4.0.5", "文本已核对"],
            ["DYNAMIXEL_XL330_M288_sample/README.md", "快速接线、安装、命令和链接", "与本手册一致"],
            ["本 DOCX", "规格、接线、控制表、实验和 MCU 迁移", "结构与备用渲染检查完成；建议在 Word 中确认自动分页"],
        ],
        [3900, 3000, 2460],
    )
    doc.add_paragraph(
        "“未在实物上运行”意味着串口、ID、供电、实际运动方向、夹具安全和 CSV 动态值仍必须由现场硬件确认；本文不把离线验证表述为整机成功。"
    )

    doc.add_heading("14. 官方链接与资料来源", level=1)
    add_source(
        doc,
        "R1",
        "ROBOTIS Docs — XL330-M288-T",
        "https://docs.robotis.com/docs/dxl/model_reference/x_series/xl_series/xl330-m288/",
        "规格、控制模式、控制表、单位、连接器、通信电平和安全警告",
    )
    add_source(
        doc,
        "R2",
        "ROBOTIS Docs — DYNAMIXEL Protocol 2.0",
        "https://docs.robotis.com/docs/dxl/protocol/protocol2/",
        "半双工 UART、包结构、方向切换、CRC 与 1.5ms 字节间隔",
    )
    add_source(
        doc,
        "R3",
        "ROBOTIS Docs — U2D2",
        "https://docs.robotis.com/docs/parts/interface/u2d2/",
        "U2D2 不向电机供电、端口、驱动和波特率",
    )
    add_source(
        doc,
        "R4",
        "ROBOTIS Docs — U2D2 Power Hub Board",
        "https://docs.robotis.com/docs/parts/interface/u2d2_power_hub/",
        "3.5–24V 板级范围、10A max、接线、极性和只能使用一个电源输入",
    )
    add_source(
        doc,
        "R5",
        "ROBOTIS Docs — DYNAMIXEL SDK",
        "https://docs.robotis.com/docs/software/dynamixel_sdk/",
        "SDK 语言、教程与下载入口",
    )
    add_source(
        doc,
        "R6",
        "ROBOTIS Docs — DYNAMIXEL Wizard 2.0",
        "https://docs.robotis.com/docs/software/dynamixel_wizard_2_0/introduction/",
        "下载、支持系统、扫描、诊断、绘图和固件功能",
    )
    add_source(
        doc,
        "R7",
        "ROBOTIS Docs — Python Basic Read/Write Tutorial",
        "https://docs.robotis.com/docs/software/dynamixel_sdk/basic_read_write_tutorial/python/",
        "PortHandler、PacketHandler、Torque、Goal/Present Position 的 Python API 模式",
    )
    add_source(
        doc,
        "R8",
        "ROBOTIS GitHub — DynamixelSDK",
        "https://github.com/ROBOTIS-GIT/DynamixelSDK",
        "官方开源 SDK、示例与 Apache-2.0 许可证",
    )
    add_source(
        doc,
        "R9",
        "ROBOTIS Japan e-Shop — XL330-M288-T",
        "https://e-shop.robotis.co.jp/product.php?id=417",
        "日本商品型号、品番 902-0163-000、同捆物和商品入口",
    )
    add_source(
        doc,
        "R10",
        "PyPI — dynamixel-sdk",
        "https://pypi.org/project/dynamixel-sdk/",
        "本样例 requirements 使用 4.0.5（2026-05-06 发布）",
    )

    doc.add_heading("附录 A：快速检查卡", level=1)
    add_table(
        doc,
        ["上电前", "上电后", "运动前", "结束后"],
        [
            ["电源 5.0V；极性正确", "Wizard 识别 1200", "电机固定、周围清空", "Torque=0"],
            ["只用一个 Power Hub 输入", "电压约 5.0V", "Goal 接近 Present", "CSV 已保存"],
            ["TTL 3Pin，不接 RS-485", "Hardware Error=0", "Profile 低速、小角度", "Power Hub 断电再拔线"],
            ["物理急停可触及", "温度正常", "Wizard 已关闭", "记录异常与参数"],
        ],
        [2340, 2340, 2340, 2340],
    )
    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
